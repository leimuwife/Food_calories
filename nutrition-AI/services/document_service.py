"""文档处理服务 - 文档解析与切片"""
import json
import os
import tempfile
from typing import List, Tuple
from loguru import logger
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from constants.global_constants import FileConstants

# 支持的文件格式（从全局常量导入）
SUPPORTED_FORMATS = FileConstants.SUPPORTED_FORMATS


class DocumentService:
    """文档处理服务"""

    def __init__(self) -> None:
        pass

    def validate_file_format(self, filename: str) -> bool:
        """
        校验文件格式是否支持

        Args:
            filename: 文件名

        Returns:
            是否支持
        """
        _, ext = os.path.splitext(filename.lower())
        return ext in SUPPORTED_FORMATS

    def parse_document(self, file_content: bytes, filename: str) -> str:
        """
        解析文档内容，提取纯文本

        Args:
            file_content: 文件二进制内容
            filename: 文件名（用于判断格式）

        Returns:
            提取的纯文本内容
        """
        ext = os.path.splitext(filename.lower())[1]
        logger.info("开始解析文档: filename={}, ext={}", filename, ext)

        # 保存到临时文件
        tmp_dir = tempfile.mkdtemp()
        tmp_path = os.path.join(tmp_dir, filename)

        try:
            with open(tmp_path, "wb") as f:
                f.write(file_content)

            # 根据格式选择解析方式
            if ext == ".pdf":
                text = self._parse_pdf(tmp_path)
            elif ext in (".txt", ".md"):
                text = self._parse_text(tmp_path)
            elif ext == ".docx":
                text = self._parse_docx(tmp_path)
            elif ext == ".doc":
                # .doc格式需要特殊处理（暂不支持，提示转换）
                raise ValueError("暂不支持.doc格式，请转换为.docx或.pdf后重试")
            elif ext in (".json", ".jsonl"):
                text = self._parse_json(file_content)
            else:
                raise ValueError(f"不支持的文件格式: {ext}")

            if not text or not text.strip():
                raise ValueError("文档内容为空")

            logger.info("文档解析成功: filename={}, text_length={}", filename, len(text))
            return text

        finally:
            # 清理临时文件
            try:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
                if os.path.exists(tmp_dir):
                    os.rmdir(tmp_dir)
            except Exception:
                pass

    def _parse_pdf(self, file_path: str) -> str:
        """解析PDF文件"""
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

        return text

    def _parse_text(self, file_path: str) -> str:
        """解析文本文件（txt/md）"""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def _parse_docx(self, file_path: str) -> str:
        """解析Word文档"""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        # 读取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"

        return text

    def _parse_json(self, file_content: bytes) -> str:
        """
        解析JSON结构化文件，将JSON数据转换为通顺的自然语言文本
        禁止直接将原始json字符串送入文本分割器，防止字段结构被切片拆分、破坏语义

        支持三种格式：
        1. 根节点为数组：循环遍历每一条结构体，拼接成通顺易懂的自然语言文本
        2. 根节点为单个json对象：格式化转为可读文本
        3. JSONL格式（每行一个JSON对象）：逐行解析后拼接
        """
        import json

        # UTF-8解码，处理BOM头
        try:
            text_raw = file_content.decode("utf-8-sig")
        except UnicodeDecodeError:
            try:
                text_raw = file_content.decode("gbk")
            except UnicodeDecodeError:
                raise ValueError("JSON文件编码无法识别，请使用UTF-8编码")

        # 尝试解析为JSONL（每行一个JSON对象）
        lines = [line.strip() for line in text_raw.strip().splitlines() if line.strip()]
        jsonl_objects = []
        is_jsonl = True
        for line in lines:
            try:
                obj = json.loads(line)
                if isinstance(obj, dict):
                    jsonl_objects.append(obj)
                else:
                    is_jsonl = False
                    break
            except json.JSONDecodeError:
                is_jsonl = False
                break

        if is_jsonl and len(jsonl_objects) > 1:
            # JSONL格式：每行一个JSON对象
            text_parts = []
            for index, item in enumerate(jsonl_objects, 1):
                text_parts.append(self._json_object_to_text(item, index))
            return "\n\n".join(text_parts)

        # 标准JSON格式
        try:
            data = json.loads(text_raw)
        except json.JSONDecodeError as e:
            raise ValueError(f"JSON格式解析失败: {str(e)}")

        # 将JSON数据转换为自然语言文本
        if isinstance(data, list):
            # 根节点为数组：遍历每一条结构体
            text_parts = []
            for index, item in enumerate(data, 1):
                if isinstance(item, dict):
                    text_parts.append(self._json_object_to_text(item, index))
                else:
                    text_parts.append(f"第{index}条记录：{str(item)}")
            return "\n\n".join(text_parts)

        elif isinstance(data, dict):
            # 根节点为单个json对象
            return self._json_object_to_text(data)

        else:
            # 简单类型（字符串/数字等）
            return str(data)

    def _json_object_to_text(self, obj: dict, index: int = None) -> str:
        """
        将单个JSON对象转换为通顺可读的自然语言文本

        去除序号前缀等模板噪声，让核心内容占据语义主导地位，
        提升向量检索的召回准确率。

        Args:
            obj: JSON字典
            index: 记录序号（保留参数兼容性，不再拼入文本）

        Returns:
            自然语言文本
        """
        # 食物数据：food_name + calorie → 自然语句，去除模板噪声
        food_name = obj.get("food_name")
        calorie = obj.get("calorie")
        if food_name and calorie is not None:
            return f"{food_name}，每100克热量{calorie}千卡。"

        # 通用处理：键值对拼接为自然语句（不再添加序号前缀）
        parts = []
        for key, value in obj.items():
            readable_key = key.replace("_", " ").replace("-", " ")
            if isinstance(value, (dict, list)):
                value_str = json.dumps(value, ensure_ascii=False)
            else:
                value_str = str(value)
            parts.append(f"{readable_key}：{value_str}")

        return "；".join(parts) + "。"

    def split_document(self, text: str, filename: str) -> List[Document]:
        """
        将文档文本切片

        Args:
            text: 文档纯文本
            filename: 文件名（用于元数据）

        Returns:
            切片后的文档列表
        """
        from config.settings import settings

        logger.info("开始文本切片: filename={}, chunk_size={}, chunk_overlap={}",
                    filename, settings.chunk_size, settings.chunk_overlap)

        # 初始化分割器
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            separators=settings.separators,
            length_function=len
        )

        # 创建文档对象
        document = Document(
            page_content=text,
            metadata={
                "filename": filename,
                "chunk_size": settings.chunk_size
            }
        )

        # 执行切片
        chunks = splitter.split_documents([document])

        logger.info("文本切片完成: filename={}, chunks={}", filename, len(chunks))
        return chunks

    def process_document(self, file_content: bytes, filename: str) -> Tuple[List[Document], int]:
        """
        完整处理流程：解析 -> 切片

        JSON/JSONL文件特殊处理：每条记录独立成一个Document，不走字符切片，
        确保单条食物数据（如"牛蛙"）能被精确召回。

        Args:
            file_content: 文件二进制内容
            filename: 文件名

        Returns:
            (切片文档列表, 切片数量)
        """
        # 校验格式
        if not self.validate_file_format(filename):
            supported = ", ".join(SUPPORTED_FORMATS)
            raise ValueError(f"不支持的文件格式，支持的格式: {supported}")

        ext = os.path.splitext(filename.lower())[1]

        # JSON/JSONL文件：每条记录独立成片，不走RecursiveCharacterTextSplitter
        if ext in (".json", ".jsonl"):
            chunks = self._parse_json_to_documents(file_content, filename)
            logger.info("JSON/JSONL独立切片完成: filename={}, records={}", filename, len(chunks))
            return chunks, len(chunks)

        # 其他格式：解析 -> 字符切片
        text = self.parse_document(file_content, filename)
        chunks = self.split_document(text, filename)

        return chunks, len(chunks)

    @staticmethod
    def _create_food_document(obj: dict, filename: str, index: int) -> Document:
        """
        创建食物数据Document

        关键设计：page_content仅存食物名（用于embedding，最大化名称权重），
        完整营养信息存入metadata["nutrition_text"]（用于检索返回展示）。
        这样向量只编码食物名称，不被"每100克热量"等模板噪声稀释。
        """
        food_name = str(obj.get("food_name"))
        calorie = obj.get("calorie")
        return Document(
            page_content=food_name,
            metadata={
                "filename": filename,
                "chunk_index": index,
                "nutrition_text": f"{food_name}，每100克热量{calorie}千卡。"
            }
        )

    @staticmethod
    def _is_food_record(obj: dict) -> bool:
        """判断JSON对象是否为食物数据（同时含food_name和calorie）"""
        return bool(obj.get("food_name")) and obj.get("calorie") is not None

    def _parse_json_to_documents(self, file_content: bytes, filename: str) -> List[Document]:
        """
        将JSON/JSONL文件中每条记录解析为独立的Document

        JSONL：每行一个JSON对象，每行一个Document
        JSON数组：每个元素一个Document
        JSON对象：单个Document

        食物数据（含food_name+calorie）特殊处理：
        page_content仅存食物名用于embedding，完整营养信息存入metadata。

        Args:
            file_content: 文件二进制内容
            filename: 文件名

        Returns:
            Document列表，每条记录一个Document
        """
        # 解码
        try:
            text_raw = file_content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text_raw = file_content.decode("gbk")
            except UnicodeDecodeError:
                raise ValueError("JSON文件编码无法识别，请使用UTF-8编码")

        ext = os.path.splitext(filename.lower())[1]
        documents = []

        if ext == ".jsonl":
            # JSONL：逐行解析，每行一个Document
            lines = [line.strip() for line in text_raw.strip().splitlines() if line.strip()]
            for index, line in enumerate(lines, 1):
                try:
                    obj = json.loads(line)
                    if isinstance(obj, dict):
                        if self._is_food_record(obj):
                            documents.append(self._create_food_document(obj, filename, index))
                        else:
                            text = self._json_object_to_text(obj, index)
                            documents.append(Document(page_content=text, metadata={
                                "filename": filename, "chunk_index": index
                            }))
                except json.JSONDecodeError:
                    logger.warning("JSONL第{}行解析失败，跳过: {}", index, line[:80])
        else:
            # 标准JSON
            try:
                data = json.loads(text_raw)
            except json.JSONDecodeError as e:
                raise ValueError(f"JSON格式解析失败: {str(e)}")

            if isinstance(data, list):
                for index, item in enumerate(data, 1):
                    if isinstance(item, dict):
                        if self._is_food_record(item):
                            documents.append(self._create_food_document(item, filename, index))
                        else:
                            text = self._json_object_to_text(item, index)
                            documents.append(Document(page_content=text, metadata={
                                "filename": filename, "chunk_index": index
                            }))
            elif isinstance(data, dict):
                if self._is_food_record(data):
                    documents.append(self._create_food_document(data, filename, 1))
                else:
                    text = self._json_object_to_text(data)
                    documents.append(Document(page_content=text, metadata={
                        "filename": filename, "chunk_index": 1
                    }))

        logger.info("JSON解析为独立文档: filename={}, records={}", filename, len(documents))
        return documents
